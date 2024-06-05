import pandas as pd

from docx import Document



# Load the CSV data

csv_file_path = "ABCBM_2024_2024_04_29.csv"

data = pd.read_csv(csv_file_path)



# Create a new Document
doc = Document()



# Define a function to add formatted data to the document

def add_record_to_doc(doc, row):

    death_registration_no = row.get('Death Reg Number', 'N/A')
    death_pre_reg_identifier = row.get('Death Pre Reg Identifier','N/A')
    name = f"{row.get('Deceased Last Name','N/A')},{row.get('Deceased Given Name','N/A')}"
    sex = row.get('Gender','N/A')
    address = row.get('Street Address of Death','N/A')
    death_date = row.get('Date of Death','N/A')
    birth_date = row.get('Birth Date','N/A')
    marital_status = row.get('Marital Type Code','N/A')
    spouse = f"{row.get('Spouse Last Name','N/A')},{row.get('Spouse Given Name','N/A')}"
    father = f"{row.get('Father Last Name','N/A')},{row.get('Father Given Name','N/A')}"
    birthplace = row.get('Deceased Birth City','N/A')
    mother = f"{row.get('Mother Maiden Last Name','N/A')},{row.get('Mother Given Name','N/A')}"
    occupation = row.get('Occupation','N/A')
    ah_phn = row.get('Personal Health Number','N/A')
    autopsy = row.get('Autopsy Performed','N/A')
    death_place = row.get('Location Code','N/A')
    underlying_cause = row.get('Underlying Cause of Death','N/A')



    # Add the formatted data to the document

    doc.add_paragraph("EXTRACT FROM DEATH CERTIFICATE\n")
    doc.add_paragraph(f"30 April 2024                              Death Registration No.{death_registration_no}\n")
    doc.add_paragraph(f"                                                 Death Pre Registration Identifier  {death_pre_reg_identifier}\n")
    doc.add_paragraph(f"Name:        {name} Sex: {sex}\n")
    doc.add_paragraph(f"Address:    {address}\n")
    doc.add_paragraph(f"Death Date:  {death_date} Birth Date:    {birth_date}\n")
    doc.add_paragraph(f"Mar. Status:{marital_status}                   Spouse:{spouse}\n")
    doc.add_paragraph(f"Father:      {father}\n")
    doc.add_paragraph(f"Birthplace:  {birthplace}\n")
    doc.add_paragraph(f"Mother:      {mother}\n")
    doc.add_paragraph(f"Occupation:  {occupation}AH_PHN  :      {ah_phn}\n")
    doc.add_paragraph(f"Autopsy:    {autopsy}\n")
    doc.add_paragraph(f"Death Place:{death_place}\n")
    doc.add_paragraph(f"Death Causes:\n")
    doc.add_paragraph(f"    Underlying Cause:{underlying_cause}\n")
    doc.add_paragraph("                        ********************************")
    doc.add_paragraph("                        *                              *")
    doc.add_paragraph("                        *      D O  N O T  C O P Y     *")
    doc.add_paragraph("                        *                              *")
    doc.add_paragraph("                        ********************************\n")
    doc.add_paragraph("      --------------------------------------------------------------------------")
    doc.add_paragraph("      CANCER CLINIC RECORD:       Region:  2 S. Alberta-JACC\n")
    doc.add_paragraph(f"      Name:        {name.lower()}Sex: {sex}")
    doc.add_paragraph(f"      Address:    {address}")
    doc.add_paragraph(f"      DIED:        {death_date}c          BORN:   {birth_date}  c        Vitstat: ?")
    doc.add_paragraph(f"      marstat:    {marital_status}                   spouse:{spouse}")
    doc.add_paragraph(f"      Birthname:{name}")
    doc.add_paragraph(f"      Birthplace:  {birthplace}AHW ULI No:        {ah_phn}")
    doc.add_paragraph(f"      Conf srce:   unofficial          CR dth_cert:  {death_registration_no}")
    doc.add_paragraph(f"      CR ICDs:    {underlying_cause}                               No. Prims:  1")
    doc.add_paragraph(f"      VS codes:    brthplc-  {birthplace}hosp-  {death_place} dthplc-  {death_place}causes:    {underlying_cause}")
    doc.add_paragraph(f"      file_loc: l       Media: l-chart                      auto upd: YES  /2\n")



# Process each row and add to the document

for index, row in data.iterrows():

    add_record_to_doc(doc,row)



# Save the document

output_file_path = "output2.docx"

doc.save(output_file_path)



print(f"Document saved to{output_file_path}")



